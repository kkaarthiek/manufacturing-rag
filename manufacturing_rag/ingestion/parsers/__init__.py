"""
Step 2-3 — format-aware parsing (spec 6.2).  STATUS: IMPLEMENTED (text extraction).

One extractor per format; binaries (pdf/xlsx/docx) are real OOXML/PDF parsers.
The redundancy rule (2 parsers + vision reconcile for binaries) is a Tranche-B
add (needs the vision LLM); here the single deterministic parser is exercised and
its output verified against the ground-truth gate.

Each `parse(fmt, raw)` returns plain text (rows joined for tabular formats).
Encoding detection (utf-8-sig/utf-8/cp1252/latin-1) handles the BOM and the
cp1252 email trap.
"""

from __future__ import annotations

import csv as _csv
import html as _html
import io
import json
import re
import zipfile

REDUNDANT_FORMATS = {"pdf", "xlsx", "docx"}  # Tranche B: 2nd (vision) parser + reconcile


# --------------------------------------------------------------------------- #
# Encoding / detection
# --------------------------------------------------------------------------- #
def decode_text(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            s = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    else:
        s = data.decode("utf-8", "replace")
    return s.replace("\r\n", "\n").replace("\r", "\n")


def detect_format(filename: str, data: bytes) -> str:
    name = filename.lower()
    if data[:4] == b"%PDF":
        return "pdf"
    if data[:2] == b"PK":
        if name.endswith(".docx"):
            return "docx"
        if name.endswith(".xlsx"):
            return "xlsx"
        return "zip"
    ext = name.rsplit(".", 1)[-1] if "." in name else ""
    return {"csv": "csv", "tsv": "csv", "html": "html", "htm": "html",
            "json": "json", "eml": "eml", "msg": "eml"}.get(ext, "txt")


# --------------------------------------------------------------------------- #
# Per-format extractors
# --------------------------------------------------------------------------- #
def _strip_html(s: str) -> str:
    s = re.sub(r"(?is)<(script|style|nav|footer)[^>]*>.*?</\1>", " ", s)
    s = re.sub(r"(?i)<li[^>]*>", "\n- ", s)
    s = re.sub(r"(?i)<(br|/p|/h[1-6]|/tr|/li|/div|/section)[^>]*>", "\n", s)
    s = re.sub(r"<[^>]+>", " ", s)
    s = _html.unescape(s)
    return re.sub(r"[ \t]+", " ", re.sub(r"\n\s*\n+", "\n\n", s)).strip()


def extract_eml(s: str):
    """Return (subject, body) with quoted history + signature stripped."""
    headers, _, body = s.partition("\n\n")
    subject = ""
    for line in headers.splitlines():
        if line.lower().startswith("subject:"):
            subject = line.split(":", 1)[1].strip()
    keep = []
    for line in body.splitlines():
        ls = line.strip()
        if ls.startswith(">"):
            continue
        if ls in ("--", "-- "):
            break
        keep.append(line)
    return subject, "\n".join(keep).strip()


def extract_xlsx(data: bytes) -> str:
    """openpyxl (open-source) -> rows as 'cell | cell'; stdlib OOXML fallback."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        out = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = ["" if c is None else str(c) for c in row]
                if any(c.strip() for c in cells):
                    out.append(" | ".join(cells))
        return "\n".join(out).strip()
    except Exception:
        return _extract_xlsx_stdlib(data)


def _extract_xlsx_stdlib(data: bytes) -> str:
    out = []
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        shared = []
        if "xl/sharedStrings.xml" in z.namelist():
            ss = z.read("xl/sharedStrings.xml").decode("utf-8", "replace")
            shared = re.findall(r"<t[^>]*>(.*?)</t>", ss, re.S)
        for nm in sorted(n for n in z.namelist()
                         if re.match(r"xl/worksheets/sheet\d+\.xml", n)):
            xml = z.read(nm).decode("utf-8", "replace")
            for row in re.findall(r"<row[^>]*>(.*?)</row>", xml, re.S):
                cells = []
                for cm in re.finditer(r"<c\b([^>]*)>(.*?)</c>", row, re.S):
                    attrs, inner = cm.group(1), cm.group(2)
                    tmatch = re.search(r't="(\w+)"', attrs)
                    ctype = tmatch.group(1) if tmatch else None
                    if ctype == "s":
                        idx = re.search(r"<v>(\d+)</v>", inner)
                        if idx:
                            cells.append(_html.unescape(shared[int(idx.group(1))]))
                    elif ctype == "inlineStr" or "<is>" in inner:
                        t = re.search(r"<t[^>]*>(.*?)</t>", inner, re.S)
                        cells.append(_html.unescape(t.group(1)) if t else "")
                    else:
                        v = re.search(r"<v>(.*?)</v>", inner, re.S)
                        cells.append(_html.unescape(v.group(1)) if v else "")
                if any(x.strip() for x in cells):
                    out.append(" | ".join(cells))
    return "\n".join(out).strip()


def extract_docx(data: bytes) -> str:
    """python-docx (open-source) -> paragraphs + table cells; stdlib fallback."""
    try:
        import docx
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
    except Exception:
        return _extract_docx_stdlib(data)


def _extract_docx_stdlib(data: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    paras = []
    for p in re.findall(r"<w:p\b.*?</w:p>", xml, re.S):
        texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", p, re.S)
        line = _html.unescape("".join(texts)).strip()
        if line:
            paras.append(line)
    return "\n".join(paras).strip()


def _pdf_decode(s: bytes) -> str:
    """Decode a PDF literal string body: handle \\n \\t \\( \\) octal \\ddd etc."""
    out = bytearray()
    i, n = 0, len(s)
    while i < n:
        c = s[i]
        if c == 0x5C and i + 1 < n:                    # backslash escape
            nx = s[i + 1]
            mp = {0x6E: 0x0A, 0x72: 0x0D, 0x74: 0x09, 0x62: 0x08, 0x66: 0x0C}
            if nx in mp:
                out.append(mp[nx]); i += 2
            elif nx in (0x28, 0x29, 0x5C):             # ( ) \
                out.append(nx); i += 2
            elif 0x30 <= nx <= 0x37:                    # octal \ddd
                od = chr(nx); i += 2
                for _ in range(2):
                    if i < n and 0x30 <= s[i] <= 0x37:
                        od += chr(s[i]); i += 1
                out.append(int(od, 8) & 0xFF)
            else:
                i += 1                                  # line continuation: drop backslash
        else:
            out.append(c); i += 1
    return out.decode("latin-1", "replace")


def _tj_array(seg: bytes) -> str:
    """Text from a [ (str) kern (str) ... ] TJ array; large -ve kern => word space."""
    parts = []
    for tok in re.finditer(rb"\(((?:\\.|[^()\\])*)\)|(-?\d+\.?\d*)", seg):
        if tok.group(1) is not None:
            parts.append(_pdf_decode(tok.group(1)))
        elif tok.group(2) and float(tok.group(2)) < -100:
            parts.append(" ")                          # kerning-induced word break
    return "".join(parts)


def extract_pdf(data: bytes) -> str:
    """Open-source PDF text extraction: PyMuPDF (fitz) primary, then pdfplumber,
    then the stdlib regex fallback (so the offline gate never hard-depends on a lib)."""
    # 1) PyMuPDF (fitz) — fast, handles real PDFs (TJ arrays, CID fonts, object streams)
    try:
        import fitz
        with fitz.open(stream=data, filetype="pdf") as doc:
            txt = "\n".join(page.get_text("text") for page in doc).strip()
        if txt:
            return txt
    except Exception:
        pass
    # 2) pdfplumber — good for table-heavy pages
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            txt = "\n".join((p.extract_text() or "") for p in pdf.pages).strip()
        if txt:
            return txt
    except Exception:
        pass
    # 3) stdlib regex fallback
    return _extract_pdf_stdlib(data)


def _extract_pdf_stdlib(data: bytes) -> str:
    """Stdlib regex PDF extractor (fallback): FlateDecode streams, (..)Tj + [..]TJ."""
    import zlib
    pages = []
    for m in re.finditer(rb"stream\r?\n?(.*?)\r?\nendstream", data, re.S):
        s = m.group(1)
        try:
            s = zlib.decompress(s)
        except Exception:
            pass
        if b"TJ" not in s and b"Tj" not in s:
            continue                                    # not a text content stream
        parts = []
        # show operators + line-positioning, in document order
        for op in re.finditer(
                rb"(\[(?:[^\[\]]|\\.)*\]\s*TJ)|(\((?:\\.|[^()\\])*\)\s*(?:Tj|'|\"))"
                rb"|(Td|TD|T\*)", s):
            if op.group(1):                             # [..] TJ
                parts.append(_tj_array(op.group(1)))
            elif op.group(2):                           # (..) Tj/'/"
                inner = re.match(rb"\(((?:\\.|[^()\\])*)\)", op.group(2))
                parts.append(_pdf_decode(inner.group(1)) if inner else "")
            else:                                       # positioning -> newline
                parts.append("\n")
        txt = re.sub(r"[ \t]+", " ", " ".join(parts))
        if txt.strip():
            pages.append(txt.strip())
    return re.sub(r"\n\s*\n+", "\n\n", "\n".join(pages)).strip()


def _rows_to_markdown(rows: list[list[str]]) -> str:
    rows = [[(c or "").replace("\n", " ").strip() for c in r] for r in rows if any(r)]
    if not rows:
        return ""
    w = max(len(r) for r in rows)
    rows = [r + [""] * (w - len(r)) for r in rows]
    out = ["| " + " | ".join(rows[0]) + " |", "|" + "|".join(["---"] * w) + "|"]
    for r in rows[1:]:
        out.append("| " + " | ".join(r) + " |")
    return "\n".join(out)


def extract_pdf_tables(data: bytes) -> list[dict]:
    """Structured tables via pdfplumber -> [{page, rows, markdown}]. Empty if none."""
    try:
        import pdfplumber
    except Exception:
        return []
    out = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for pi, page in enumerate(pdf.pages):
                for tbl in (page.extract_tables() or []):
                    rows = [[(c or "").strip() for c in row] for row in tbl if any(row)]
                    if len(rows) >= 2 and len(rows[0]) >= 2:
                        out.append({"page": pi, "rows": rows,
                                    "markdown": _rows_to_markdown(rows)})
    except Exception:
        return out
    return out


def pdf_text_coverage(data: bytes) -> float:
    """Avg text chars per page. Low (< ~80) => likely scanned/image-only -> OCR."""
    try:
        import fitz
        with fitz.open(stream=data, filetype="pdf") as doc:
            if doc.page_count == 0:
                return 0.0
            chars = sum(len(p.get_text()) for p in doc)
            return chars / doc.page_count
    except Exception:
        return 1e9  # unknown -> assume has text (don't force OCR)


def extract_pdf_images(data: bytes, max_images: int = 12) -> list[tuple]:
    """Embedded raster images -> [(png_bytes, page_index)] (for vision captioning)."""
    try:
        import fitz
    except Exception:
        return []
    out = []
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            for pi in range(doc.page_count):
                for img in doc[pi].get_images(full=True):
                    try:
                        pix = fitz.Pixmap(doc, img[0])
                        if pix.n - pix.alpha > 3:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        if pix.width >= 64 and pix.height >= 64:   # skip tiny icons
                            out.append((pix.tobytes("png"), pi))
                    except Exception:
                        pass
                    if len(out) >= max_images:
                        return out
    except Exception:
        return out
    return out


def render_pdf_pages(data: bytes, max_pages: int = 8, dpi: int = 150) -> list[tuple]:
    """Render pages to PNG -> [(png_bytes, page_index)] for vision OCR of scans."""
    try:
        import fitz
    except Exception:
        return []
    out = []
    try:
        with fitz.open(stream=data, filetype="pdf") as doc:
            for pi in range(min(doc.page_count, max_pages)):
                pix = doc[pi].get_pixmap(dpi=dpi)
                out.append((pix.tobytes("png"), pi))
    except Exception:
        return out
    return out


def parse(fmt: str, raw: bytes):
    """Return text for the given format (text formats decoded; binaries parsed)."""
    if fmt == "pdf":
        return extract_pdf(raw)
    if fmt == "xlsx":
        return extract_xlsx(raw)
    if fmt == "docx":
        return extract_docx(raw)
    s = decode_text(raw)
    if fmt == "html":
        return _strip_html(s)
    if fmt == "eml":
        subj, body = extract_eml(s)
        return f"{subj}\n{body}" if subj else body
    if fmt == "json":
        try:
            return json.dumps(json.loads(s), indent=2, ensure_ascii=False)
        except Exception:
            return s
    if fmt == "csv":
        rows = list(_csv.reader(io.StringIO(s)))
        return "\n".join(" | ".join(r) for r in rows)
    return s


def parse_csv_rows(raw: bytes):
    """CSV-aware: return list[dict] keyed by normalized header (for assemblers)."""
    s = decode_text(raw)
    reader = _csv.reader(io.StringIO(s))
    rows = list(reader)
    if not rows:
        return []
    header = [h.strip() for h in rows[0]]
    out = []
    for r in rows[1:]:
        if not any(c.strip() for c in r):
            continue
        out.append({header[i]: (r[i].strip() if i < len(r) else "")
                    for i in range(len(header))})
    return out


__all__ = ["parse", "parse_csv_rows", "detect_format", "decode_text",
           "extract_eml", "extract_xlsx", "extract_docx", "extract_pdf",
           "REDUNDANT_FORMATS"]
